#Import core python modules
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

#Low-Level imports to make URLs work
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

#For timestamping the file_name
from datetime import datetime
import re

# Function to remove XML-incompatible characters
def sanitize_for_xml(text):
    if not text:
        return ""
    # Remove control characters that aren't allowed in XML (except tab, newline, carriage return)
    # Characters < 32 except 9, 10, 13
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', str(text))

#Function that adds clickable links to the document
def add_hyperlink(paragraph, url, text):

    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1') 
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rPr.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')
    rPr.append(sz)
    
    new_run.append(rPr)
    
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

#Function that creates the actual document
def create_citations(private_data, public_data):
    
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    #Stylized Heading 1
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.name = 'Arial'
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True
    heading1_style.font.color.rgb = RGBColor(0, 0, 0)

    #Stylized Heading 2
    heading2_style = doc.styles['Heading 2']
    heading2_style.font.name = 'Arial'
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True
    heading2_style.font.color.rgb = RGBColor(0, 0, 0)
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    title = doc.add_heading("Citations & Source Document", level=1)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    
    heading = doc.add_heading("Public Sources", level=2)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)

    # Extract URLs from public_data dict
    source_urls = public_data.get("source_urls", [])
    citations = public_data.get("citations", [])
    
    # 1. List Public Sources
    if source_urls:
        intro = doc.add_paragraph(f"{len(source_urls)} public web pages reviewed:")
        intro.paragraph_format.space_after = Pt(6)

        for idx, url in enumerate(source_urls, start=1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            
            run = p.add_run(f"{idx}. ")
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            
            add_hyperlink(p, url, url)
    else:
        # Fallback if no sources exist
        doc.add_paragraph("No public web sources used")

    # 2. Detailed Citations Table
    if citations:
        doc.add_heading("Key Claims & Sources", level=2)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Claim / Fact'
        hdr_cells[1].text = 'Source'
        
        # Populate
        for item in citations:
            claim_text = sanitize_for_xml(item.get("claim", ""))
            source_text = sanitize_for_xml(item.get("source", ""))
            
            row_cells = table.add_row().cells
            row_cells[0].text = claim_text
            row_cells[1].text = source_text
        
     # Timestamped output filename
     # Timestamped output filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Calculate project root (assuming this script is in scripts/ subdir)
    import os
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    
    output_dir = os.path.join(project_root, "output", "citations")
    os.makedirs(output_dir, exist_ok=True)
    
    output_filename = os.path.join(output_dir, f"company_citations_{timestamp}.docx")
    
    doc.save(output_filename)

